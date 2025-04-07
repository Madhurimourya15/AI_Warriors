# RAG-based Multimodal Document Processor with OpenAI/Gemini and Security Guardrails

import streamlit as st
import tempfile
import os
import base64
import ssl
import json
import pandas as pd
import re
from docx import Document as DocxDocument
from PIL import Image
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader, CSVLoader, UnstructuredExcelLoader, UnstructuredPowerPointLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
import google.generativeai as genai
import openai
import base64
from io import BytesIO

# ===================== SSL Fix ==========================
try:
    ssl._create_default_https_context = ssl._create_unverified_context
except Exception:
    pass

# ===================== Helper Functions ==========================
def get_base64_download_link(content, filename, mime):
    b64 = base64.b64encode(content.encode() if isinstance(content, str) else content).decode()
    return f'<a href="data:{mime};base64,{b64}" download="{filename}">{filename}</a>'

def to_json(summary):
    return json.dumps({"summary": summary}, indent=2)

def to_xml(summary):
    esc = summary.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return f"<summary>\n{esc}\n</summary>"

def to_excel(summary):
    df = pd.DataFrame([[summary]], columns=["Summary"])
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        df.to_excel(tmp.name, index=False)
        with open(tmp.name, "rb") as f:
            return f.read()

def to_word(summary):
    doc = DocxDocument()
    doc.add_paragraph(summary)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        with open(tmp.name, "rb") as f:
            return f.read()

def to_pdf(text):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        c = canvas.Canvas(tmp.name, pagesize=A4)
        width, height = A4
        text_object = c.beginText(20 * mm, height - 20 * mm)
        text_object.setFont("Helvetica", 10)
        lines = text.split("\n")
        for line in lines:
            while len(line) > 120:
                text_object.textLine(line[:120])
                line = line[120:]
            text_object.textLine(line)
        c.drawText(text_object)
        c.showPage()
        c.save()
        with open(tmp.name, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        return f"data:application/pdf;base64,{b64}"

# ===================== Security Guardrails ==========================
def is_valid_file(uploaded_file):
    if uploaded_file.size > 10 * 1024 * 1024:
        st.error("❌ File too large. Limit is 10MB.")
        return False
    if any(char in uploaded_file.name for char in [';', '|', '$']):
        st.error("❌ Invalid file name.")
        return False
    return True

def redact_pii(text):
    text = re.sub(r'\b\d{3}-\d{2}-\d{4}\b', '[REDACTED SSN]', text)
    text = re.sub(r'\b\d{10}\b', '[REDACTED PHONE]', text)
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[REDACTED EMAIL]', text)
    return text

def filter_output(output):
    banned_keywords = ['password', 'confidential', 'classified']
    for word in banned_keywords:
        if word.lower() in output.lower():
            return "[FILTERED CONTENT DETECTED]"
    return output

# ===================== Main App ==========================
def main():
    st.set_page_config(page_title="Document Processor Using GenAI")
    # 🖼️ Large logo on the top-left with title next to it
    st.image("tech.png", width=200)  # Change width or path as needed
    ##################################################################
    st.title("📑 Document Processor Using AI")
    

    provider = st.selectbox("🤖 Choose Model Provider", ["OpenAI", "Gemini"])
    api_key = st.text_input("🔐 Enter your API Key", type="password")

    if provider == "OpenAI":
        openai_model = st.selectbox("🧠 Select OpenAI Model", ["gpt-3.5-turbo", "gpt-4", "gpt-4o"])
    else:
        # temperature = st.slider("🌡️ Select temperature (Gemini only)", 0.0, 1.0, 0.7)
        temperature = st.number_input("🌡️ Enter temperature for summary generation", min_value=0.1, max_value=1.0, value=0.7, step=0.1)
    
    task = st.selectbox("🧠 Select Task", ["Document Summarization", "Named Entity Extraction", "Data Extraction"])
    uploaded_file = st.file_uploader("📤 Upload a document", type=["pdf", "docx", "txt", "csv", "xlsx", "pptx", "jpg", "jpeg", "png"])

    if st.button("🚀 Submit"):
        if not api_key:
            st.warning("Please enter your API key.")
            return

        documents = []
        if uploaded_file and is_valid_file(uploaded_file):
            suffix = uploaded_file.name.lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix="." + suffix.split(".")[-1]) as tmp:
                tmp.write(uploaded_file.getvalue())
                file_path = tmp.name

            if suffix.endswith("pdf"):
                loader = PyPDFLoader(file_path)
                documents = loader.load_and_split()
            elif suffix.endswith("docx"):
                loader = Docx2txtLoader(file_path)
                documents = loader.load_and_split()
            elif suffix.endswith("txt"):
                loader = TextLoader(file_path)
                documents = loader.load_and_split()
            elif suffix.endswith(("xlsx", "xls")):
                loader = UnstructuredExcelLoader(file_path)
                documents = loader.load_and_split()
            elif suffix.endswith("csv"):
                loader = CSVLoader(file_path)
                documents = loader.load_and_split()
            elif suffix.endswith("pptx"):
                loader = UnstructuredPowerPointLoader(file_path)
                documents = loader.load_and_split()
            elif suffix.endswith(("jpg", "jpeg", "png")):
                image = Image.open(file_path)
                buffered = BytesIO()
                image.save(buffered, format="PNG")
                b64_image = base64.b64encode(buffered.getvalue()).decode()

                if provider == "OpenAI" and openai_model in ["gpt-4", "gpt-4o"]:
                    openai.api_key = api_key
                    response = openai.ChatCompletion.create(
                        model=openai_model,
                        messages=[
                            {"role": "user", "content": [
                                {"type": "text", "text": "Extract all readable text from this image."},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64_image}"}}
                            ]}
                        ],
                        max_tokens=2048
                    )
                    documents = [Document(page_content=response.choices[0].message.content)]
                else:
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel("gemini-1.5-pro", generation_config={
                        "temperature": temperature,
                        "top_p": 1,
                        "top_k": 1,
                        "max_output_tokens": 2048
                    })
                    response = model.generate_content(["Extract all readable text from this image:", image])
                    documents = [Document(page_content=response.text)]

        if documents:
            splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = splitter.split_documents(documents)
            embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
            db = FAISS.from_documents(chunks, embeddings)
            retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 5})

            queries = {
                "Document Summarization": "Generate a detailed summary with title, intro, bullets, and conclusion.",
                "Named Entity Extraction": "Extract all named entities (people, places, dates, orgs, amounts).",
                "Data Extraction": "Extract structured tables and key data from this document."
            }

            with st.spinner("🔍 Processing with AI..."):
                docs = retriever.get_relevant_documents(queries[task])
                context = "\n\n".join([doc.page_content for doc in docs])
                context = redact_pii(context)
                safe_context = context.replace("{", "{{").replace("}", "}}")

                prompt = f"""
You are a secure assistant trained to follow privacy and compliance guardrails.
Only return safe, non-sensitive, and compliant results.
Use the following context:
{safe_context}

Task: {queries[task]}
Result:"""

                if provider == "OpenAI":
                    openai.api_key = api_key
                    response = openai.ChatCompletion.create(
                        model=openai_model,
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=2048
                    )
                    result = response.choices[0].message.content
                else:
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel("gemini-1.5-pro", generation_config={
                        "temperature": temperature,
                        "top_p": 1,
                        "top_k": 1,
                        "max_output_tokens": 2048
                    })
                    result = model.generate_content(prompt).text

                result = filter_output(result)
                st.subheader("📋 Result:")
                st.write(result)

                st.markdown("---")
                st.markdown(get_base64_download_link(to_json(result), "summary.json", "application/json"), unsafe_allow_html=True)
                st.markdown(get_base64_download_link(to_xml(result), "summary.xml", "application/xml"), unsafe_allow_html=True)
                st.markdown(get_base64_download_link(to_excel(result), "summary.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"), unsafe_allow_html=True)
                st.markdown(get_base64_download_link(to_word(result), "summary.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)
                pdf_data_uri = to_pdf(result)
                st.markdown(f'<a href="{pdf_data_uri}" download="summary.pdf">Download PDF</a>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
